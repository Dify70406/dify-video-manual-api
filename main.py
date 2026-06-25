import os
import time
import requests
import traceback
import subprocess
import re
import glob
import base64

from flask import Flask, request, jsonify, send_file
from google import genai
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

GEMINI_API_KEY = os.environ["GEMINI_API_KEY"]
client = genai.Client(api_key=GEMINI_API_KEY)


@app.route("/")
def hello():
    return "Video Manual API OK"


@app.route("/manual", methods=["POST"])
def manual():
    data = request.get_json(silent=True) or {}

    video_url = data.get("video_url")
    file_name = data.get("file_name")
    content_b64 = data.get("content")

    print("[manual] request received")
    print(f"[manual] video_url exists={bool(video_url)}")
    print(f"[manual] file_name={file_name}")
    print(f"[manual] content exists={content_b64 is not None}")
    print(f"[manual] content length={len(content_b64) if content_b64 else 0}")

    try:
        # ① 従来どおり video_url が来た場合
        if video_url:
            print("[manual] path=video_url")

            if is_youtube_url(video_url):
                print("[manual] detected youtube url")
                subtitle_text = get_youtube_subtitle(video_url)

                prompt = f"""
以下はYouTube動画の字幕です。
この字幕をもとに、操作手順書を作成してください。

必ず以下の形式だけで出力してください。
前置き、あいさつ、補足説明は禁止です。
「承知しました」「以下に示します」などの文章は出力しないでください。

字幕:
{subtitle_text}

出力形式:

# 概要

概要説明

# 操作手順

## 手順1
説明

## 手順2
説明

## 手順3
説明

# 注意事項

条件:
- Markdown形式で出力してください
- 上記の見出し以外は出力しないでください
- 前置きや会話文は禁止です
- 字幕から分かる内容だけを使ってください
- 推測しすぎないでください
"""

                response = client.models.generate_content(
                    model="gemini-2.5-pro",
                    contents=prompt
                )

                manual_text = normalize_manual_text(response.text)
                print(f"[manual] generated manual length={len(manual_text)}")

                return jsonify({
                    "manual": manual_text,
                    "source": "youtube_subtitle"
                })

            else:
                print("[manual] downloading video from url")
                work_id = str(int(time.time()))
                video_path = f"/tmp/{work_id}.mp4"

                r = requests.get(video_url, timeout=300)
                r.raise_for_status()

                with open(video_path, "wb") as f:
                    f.write(r.content)

                print(f"[manual] downloaded video_path={video_path}")
                print(f"[manual] downloaded video_size={os.path.getsize(video_path)}")

                manual_text = analyze_video_file(video_path)
                manual_text = normalize_manual_text(manual_text)
                print(f"[manual] generated manual length={len(manual_text)}")

                return jsonify({
                    "manual": manual_text,
                    "source": "video_url"
                })

        # ② Power Automate から Base64 動画が来た場合
        elif file_name and content_b64:
            print("[manual] path=sharepoint_file")

            ext = os.path.splitext(file_name)[1].lower()
            if not ext:
                ext = ".mp4"

            work_id = str(int(time.time()))
            video_path = f"/tmp/{work_id}{ext}"

            # data:video/...;base64,... のような形式にも対応
            if "," in content_b64:
                content_b64 = content_b64.split(",", 1)[1]

            video_bytes = base64.b64decode(content_b64)

            with open(video_path, "wb") as f:
                f.write(video_bytes)

            print(f"[manual] saved video_path={video_path}")
            print(f"[manual] saved video_size={os.path.getsize(video_path)}")

            manual_text = analyze_video_file(video_path)
            manual_text = normalize_manual_text(manual_text)
            print(f"[manual] generated manual length={len(manual_text)}")

            return jsonify({
                "manual": manual_text,
                "source": "sharepoint_file",
                "file_name": file_name
            })

        else:
            print("[manual] invalid request: missing video_url or file_name+content")
            return jsonify({
                "error": "video_url または file_name + content が必要です"
            }), 400

    except Exception as e:
        print("[manual] exception occurred")
        print(traceback.format_exc())
        return jsonify({
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


def analyze_video_file(video_path: str) -> str:
    print(f"[analyze_video_file] uploading video_path={video_path}")
    print(f"[analyze_video_file] video_size={os.path.getsize(video_path)}")

    # files.upload の簡易リトライ
    uploaded_file = None
    last_error = None

    for attempt in range(1, 4):
        try:
            print(f"[analyze_video_file] files.upload attempt={attempt}")
            uploaded_file = client.files.upload(file=video_path)
            print(f"[analyze_video_file] uploaded name={uploaded_file.name}")
            print(f"[analyze_video_file] initial state={uploaded_file.state.name}")
            break
        except Exception as e:
            last_error = e
            err_text = str(e)
            print(f"[analyze_video_file] files.upload error attempt={attempt}: {err_text}")
            print(traceback.format_exc())

            if "503" in err_text or "UNAVAILABLE" in err_text or "Service Unavailable" in err_text:
                wait_sec = attempt * 10
                print(f"[analyze_video_file] files.upload retry after {wait_sec}s")
                time.sleep(wait_sec)
                continue

            raise

    if uploaded_file is None:
        raise Exception(f"files.upload failed after retries: {last_error}")

    while uploaded_file.state.name == "PROCESSING":
        print("[analyze_video_file] file still processing...")
        time.sleep(5)
        uploaded_file = client.files.get(name=uploaded_file.name)
        print(f"[analyze_video_file] current state={uploaded_file.state.name}")

    if uploaded_file.state.name != "ACTIVE":
        raise Exception(f"Gemini video processing failed: {uploaded_file.state.name}")

    prompt = """
この動画を分析して操作手順書を作成してください。

必ず以下の形式だけで出力してください。
前置き、あいさつ、補足説明は禁止です。
「承知しました」「以下に示します」などの文章は出力しないでください。

出力形式:

# 概要

概要説明

# 操作手順

## 手順1
説明

## 手順2
説明

## 手順3
説明

# 注意事項

条件:
- Markdown形式で出力してください
- 上記の見出し以外は出力しないでください
- 前置きや会話文は禁止です
- 動画に具体的な操作手順が含まれていない場合は、その旨を簡潔に記載してください
"""

    print("[analyze_video_file] generate_content start")

    last_error = None

    for attempt in range(1, 4):
        try:
            print(f"[analyze_video_file] generate_content attempt={attempt}")

            response = client.models.generate_content(
                model="gemini-2.5-pro",
                contents=[
                    uploaded_file,
                    prompt
                ]
            )

            text = response.text if response.text else ""
            print(f"[analyze_video_file] response length={len(text)}")

            if not text.strip():
                raise Exception("Gemini returned empty text")

            return text

        except Exception as e:
            last_error = e
            err_text = str(e)
            print(f"[analyze_video_file] generate_content error attempt={attempt}: {err_text}")
            print(traceback.format_exc())

            # 503 / UNAVAILABLE 系だけ待って再試行
            if "503" in err_text or "UNAVAILABLE" in err_text or "Service Unavailable" in err_text:
                wait_sec = attempt * 10
                print(f"[analyze_video_file] retry after {wait_sec}s")
                time.sleep(wait_sec)
                continue

            raise

    raise Exception(f"generate_content failed after retries: {last_error}")


@app.route("/word", methods=["POST"])
def word():
    try:
        data = request.get_json(silent=True) or {}

        text = data.get("text", "")
        file_name = data.get("file_name")
        content_b64 = data.get("content")

        print("[word] request received")
        print(f"[word] file_name={file_name}")
        print(f"[word] text_length={len(text) if text else 0}")
        print(f"[word] content_exists={content_b64 is not None}")
        print(f"[word] content_length={len(content_b64) if content_b64 else 0}")

        if not text:
            text = "手順を生成できませんでした"

        text = normalize_manual_text(text)
        print(f"[word] normalized_text_length={len(text)}")

        doc_path = f"/tmp/manual_{int(time.time())}.docx"
        screenshot_paths = []

        # 動画が渡されていればスクリーンショットを抽出
        if file_name and content_b64:
            ext = os.path.splitext(file_name)[1].lower()
            if not ext:
                ext = ".mp4"

            work_id = str(int(time.time()))
            video_path = f"/tmp/word_src_{work_id}{ext}"

            if "," in content_b64:
                content_b64 = content_b64.split(",", 1)[1]

            video_bytes = base64.b64decode(content_b64)

            with open(video_path, "wb") as f:
                f.write(video_bytes)

            print(f"[word] saved video_path={video_path}")
            print(f"[word] video_size={os.path.getsize(video_path)}")

        try:
            screenshot_paths = extract_screenshots(
                video_path=video_path,
                work_id=work_id,
                max_shots=4,
            )
        except Exception:
                print("[word] screenshot extraction failed")
                print(traceback.format_exc())
                screenshot_paths = []

        print(f"[word] screenshot_count={len(screenshot_paths)}")
        print(f"[word] screenshot_paths={screenshot_paths}")

        doc = build_manual_doc(text, screenshot_paths)
        doc.save(doc_path)
        print(f"[word] doc saved: {doc_path}, size={os.path.getsize(doc_path)}")

        return send_file(
            doc_path,
            as_attachment=True,
            download_name="manual.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("[word] exception occurred")
        print(traceback.format_exc())
        return jsonify({
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


def build_manual_doc(text: str, screenshot_paths: list[str]) -> Document:
    print("[build_manual_doc] start")
    print(f"[build_manual_doc] screenshot_count={len(screenshot_paths)}")

    doc = Document()
    doc.add_heading("操作手順書", level=1)

    lines = [line.rstrip() for line in text.splitlines()]
    blocks = parse_manual_blocks(lines)

    print(f"[build_manual_doc] block_count={len(blocks)}")
    for i, block in enumerate(blocks, start=1):
        print(
            f"[build_manual_doc] block{i}: "
            f"type={block['type']}, title={block['title']}, body_lines={len(block['body'])}"
        )

    shot_index = 0

    for block in blocks:
        block_type = block["type"]
        block_title = block["title"]
        block_body = block["body"]

        if block_type == "h1":
            doc.add_heading(block_title, level=1)

        elif block_type == "h2":
            doc.add_heading(block_title, level=2)

            # 手順見出しの直後にスクリーンショットを挿入
            if shot_index < len(screenshot_paths):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(screenshot_paths[shot_index], width=Inches(5.5))
                shot_index += 1

        for body_line in block_body:
            body_line = body_line.strip()
            if body_line:
                doc.add_paragraph(body_line)

    # スクショが余ったら最後に参考画像として追加
    while shot_index < len(screenshot_paths):
        doc.add_paragraph("参考画像")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(screenshot_paths[shot_index], width=Inches(5.5))
        shot_index += 1

    return doc


def parse_manual_blocks(lines: list[str]) -> list:
    blocks = []
    current = None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("## "):
            if current:
                blocks.append(current)
            current = {
                "type": "h2",
                "title": line.replace("## ", "").strip(),
                "body": []
            }
        elif line.startswith("# "):
            if current:
                blocks.append(current)
            current = {
                "type": "h1",
                "title": line.replace("# ", "").strip(),
                "body": []
            }
        else:
            if current is None:
                current = {
                    "type": "text",
                    "title": "",
                    "body": []
                }
            current["body"].append(line)

    if current:
        blocks.append(current)

    return blocks


def extract_screenshots(video_path: str, work_id: str, max_shots: int = 4) -> list[str]:
    """
    動画の長さに応じて均等な位置からスクリーンショットを抽出する。
    """
    out_dir = f"/tmp/screens_{work_id}"
    os.makedirs(out_dir, exist_ok=True)

    duration = get_video_duration(video_path)
    if duration <= 0:
        print("[extract_screenshots] duration <= 0")
        return []

    print(f"[extract_screenshots] duration={duration}")

    timestamps = []
    for i in range(max_shots):
        pos = (i + 1) / (max_shots + 1)
        sec = max(0.5, duration * pos)
        timestamps.append(sec)

    print(f"[extract_screenshots] timestamps={timestamps}")

    screenshot_paths = []

    for idx, sec in enumerate(timestamps, start=1):
        out_path = os.path.join(out_dir, f"shot_{idx:02d}.jpg")

        cmd = [
            "ffmpeg",
            "-y",
            "-ss",
            str(sec),
            "-i",
            video_path,
            "-frames:v",
            "1",
            "-q:v",
            "2",
            out_path,
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
        )

        print(f"[extract_screenshots] ffmpeg idx={idx}, returncode={result.returncode}")

        if result.returncode == 0 and os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            screenshot_paths.append(out_path)
            print(f"[extract_screenshots] saved={out_path}, size={os.path.getsize(out_path)}")
        else:
            print(f"[extract_screenshots] failed idx={idx}, stderr={result.stderr}")

    return screenshot_paths


def get_video_duration(video_path: str) -> float:
    """
    ffprobe で動画秒数を取得
    """
    cmd = [
        "ffprobe",
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        video_path,
    ]

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=30,
    )

    print(f"[get_video_duration] returncode={result.returncode}")
    print(f"[get_video_duration] stdout={result.stdout}")
    print(f"[get_video_duration] stderr={result.stderr}")

    if result.returncode != 0:
        return 0.0

    try:
        return float(result.stdout.strip())
    except Exception:
        return 0.0


def normalize_manual_text(text: str) -> str:
    """
    AIが前置き文を返した場合に、
    最初の # 概要 / # 操作手順 / # 注意事項 / ## 手順1 から始まるように整形する。
    """
    if not text:
        return ""

    text = text.strip()

    # ```markdown ... ``` や ``` ... ``` を除去
    text = text.replace("```markdown", "").replace("```md", "").replace("```", "").strip()

    lines = text.splitlines()

    start_index = None
    heading_patterns = (
        "# 概要",
        "#操作手順",
        "# 操作手順",
        "# 注意事項",
        "## 手順1",
    )

    for i, line in enumerate(lines):
        line = line.strip()
        if any(line.startswith(pattern) for pattern in heading_patterns):
            start_index = i
            break

    if start_index is not None:
        text = "\n".join(lines[start_index:]).strip()

    return text


def is_youtube_url(url: str) -> bool:
    return "youtube.com" in url or "youtu.be" in url


def get_youtube_subtitle(url: str) -> str:
    cleanup_subtitle_files()

    outtmpl = "/tmp/youtube_subtitle"

    cmd = [
        "yt-dlp",
        "--skip-download",
        "--write-subs",
        "--write-auto-subs",
        "--sub-langs",
        "ja,en",
        "--sub-format",
        "vtt",
        "-o",
        outtmpl,
        url,
    ]

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        return f"""
字幕取得に失敗しました。

yt-dlp error:
{result.stderr}
"""

    files = glob.glob("/tmp/youtube_subtitle*.vtt")
    return get_subtitle_text(files)


def get_subtitle_text(files: list[str]) -> str:
    if not files:
        return "字幕ファイルが見つかりませんでした。"

    subtitle_text = ""

    for file in files:
        with open(file, "r", encoding="utf-8", errors="ignore") as f:
            subtitle_text += "\n" + vtt_to_text(f.read())

    if not subtitle_text.strip():
        return "字幕が空でした。"

    return subtitle_text


def cleanup_subtitle_files():
    for file in glob.glob("/tmp/youtube_subtitle*"):
        try:
            os.remove(file)
        except Exception:
            pass

def vtt_to_text(vtt: str) -> str:
    lines = vtt.splitlines()
    texts = []

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line == "WEBVTT":
            continue

        if "-->" in line:
            continue

        if re.match(r"^\d+$", line):
            continue

        line = re.sub(r"<[^>]+>", "", line)
        line = line.strip()

        if line:
            texts.append(line)

    return "\n".join(texts)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
